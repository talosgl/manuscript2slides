"""Tests for all chunking strategies as well as helpers that perform heading detection.

Known untested: We don't test that page break detection works properly during the other chunking strategies.
Is there an easy way to test that?
"""

import pytest
from docx import Document
from pathlib import Path
from manuscript2slides.processing import chunking


def test_chunk_by_paragraph_skips_empty(
    path_to_sample_docx_with_everything: Path,
) -> None:
    """
    Pass in a real document object that may contain empty paragraphs. Chunk it and get an initial count.
    Then add two more empty paragraphs to the original object, and chunk the mutated object again. Compare the before/after
    list lenths; they should be identical.

    In other words, Knowing that you have N non-empty paragraphs, add M empty ones, and assert you get exactly N chunks.
    """
    sample_document = Document(str(path_to_sample_docx_with_everything))
    initial_chunk_list = chunking.chunk_by_paragraph(sample_document)

    # Make sure this didn't bring in an empty doc somehow
    assert len(initial_chunk_list) > 0

    # Add two more (empty) paragraphs
    sample_document.add_paragraph()
    sample_document.add_paragraph()

    after_chunk_list = chunking.chunk_by_paragraph(sample_document)

    assert len(initial_chunk_list) == len(after_chunk_list)


def test_chunk_by_page_respects_page_breaks(
    path_to_sample_docx_with_everything: Path,
) -> None:
    """
    Pass in a real document with known page breaks. Verify page breaks are detected by python-docx.
    Run chunk_by_page and verify at least 1 chunk is created.
    """
    sample_document = Document(str(path_to_sample_docx_with_everything))
    # Verify the test document actually has page breaks
    page_break_count = sum(
        1 for para in sample_document.paragraphs if para.contains_page_break
    )
    assert page_break_count > 0, (
        "Test document should have at least one rendered page break"
    )

    # Run chunking and verify it produces chunks
    chunk_list = chunking.chunk_by_page(sample_document)
    assert len(chunk_list) > 0


def test_chunk_by_heading_flat_creates_slide_per_heading(
    path_to_empty_docx: Path,
) -> None:
    """Create a blank document and populate it with a known count of headings and text.
    Then verify the chunk count matches the heading count."""
    # Arrange: Create a blank document and populate it with a known count of headings and text.
    new_docx = Document(str(path_to_empty_docx))

    new_docx.add_heading("Add a Heading 1", level=1)
    new_docx.add_paragraph("Normal text under H1")
    new_docx.add_paragraph("More text under H1")

    new_docx.add_heading("And add Heading 2", level=2)
    new_docx.add_paragraph("Normal text under H2")

    new_docx.add_heading("Heading 1 again", level=1)
    new_docx.add_paragraph("Again Normal text under H1")

    # Action: chunk by heading flat
    chunk_list = chunking.chunk_by_heading_flat(new_docx)

    # There were 3 headings with no page breaks, so there should be 3 chunks
    assert len(chunk_list) == 3


def test_chunk_by_heading_nested_successfully_nests(
    path_to_empty_docx: Path,
) -> None:
    """Create a blank document and populate it with a H1, H2, and H1 structure. Then verify
    we only end up with 2 chunks."""
    # Arrange: Create a blank document and populate it with a known count of headings and text.
    new_docx = Document(str(path_to_empty_docx))

    new_docx.add_heading("Add a Heading 1", level=1)
    new_docx.add_paragraph("Normal text under H1")
    new_docx.add_paragraph("More text under H1")

    new_docx.add_heading("And add Heading 2", level=2)
    new_docx.add_paragraph("Normal text under H2")

    new_docx.add_heading("Heading 1 again", level=1)
    new_docx.add_paragraph("Again Normal text under H1")

    # Action: chunk by heading nested
    chunk_list = chunking.chunk_by_heading_nested(new_docx)

    # There were 3 headings with no page breaks H1->H2->H1, so there should be 2 chunks (H1+H2, H1)
    assert len(chunk_list) == 2


@pytest.mark.parametrize(
    "input_str,expected",
    [
        # True values
        ("Heading 1", True),
        ("Heading 2", True),
        ("Heading 3", True),
        # False values
        ("Normal", False),
        ("Header 1", False),
        ("Heading9", False),  # must be style.name not style_id
    ],
)
def test_is_standard_heading(input_str: str, expected: bool) -> None:
    """Ensure standard heading detection works as expected. We only support style.name because python-docx
    does not support reliable lookup by style.style_id.

    See:
    https://python-docx.readthedocs.io/en/latest/user/styles-understanding.html#identifying-a-style
    https://python-docx.readthedocs.io/en/latest/user/styles-using.html
    https://python-docx.readthedocs.io/en/latest/_modules/docx/styles/styles.html"""
    assert chunking.is_standard_heading(input_str) == expected


@pytest.mark.parametrize(
    "input_str,output_num",
    [
        # valid headings
        ("Heading 1", 1),
        ("Heading 9", 9),
        ("Heading 11", 11),
        # invalid headings
        ("Normal", float("inf")),
        ("Heading9", float("inf")),  # must be style.name not style_id
    ],
)
def test_get_heading_level(input_str: str, output_num: int | float) -> None:
    """Ensure we return the output number expected for standard headings, and infinity for those not supported."""
    assert chunking.get_heading_level(input_str) == output_num
