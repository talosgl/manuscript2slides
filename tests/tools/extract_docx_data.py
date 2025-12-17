"""
Extract baseline JSON data from a docx file for integration testing.

Usage:
    python tests/tools/extract_docx_data.py

This tool uses manuscript2slides internal functions to extract comprehensive
formatting and annotation data from docx files into JSON baselines for testing.
"""

from pathlib import Path
import json
import sys

from docx import Document

from manuscript2slides.processing.formatting import (
    get_theme_fonts_from_docx_package,
    get_effective_font_name_docx,
)
from manuscript2slides.annotations.extract import (
    get_all_docx_comments,
    get_all_docx_footnotes,
    get_all_docx_endnotes,
)
from manuscript2slides.internals.define_config import UserConfig
import xml.etree.ElementTree as ET

from extraction_utils import (
    rgb_to_hex,
    safe_pprint,
    filter_none_keep_false,
    XML_NS,
)


# ============================================================================
# CONFIGURATION - Edit these to change input/output files
# ============================================================================
INPUT_DOCX = "tests/data/sample_doc.docx"
OUTPUT_JSON = "tests/baselines/docx_sample.json"
# Alternative configs:
# INPUT_DOCX = "tests/data/test_formatting.docx"
# OUTPUT_JSON = "tests/baselines/docx_formatting.json"
# ============================================================================


def main() -> None:
    """Load a test docx and explore its structure."""

    # Configuration
    test_docx_path = Path(INPUT_DOCX)

    if not test_docx_path.exists():
        print(f"Error: {test_docx_path} not found!", file=sys.stderr)
        print(f"Please check the INPUT_DOCX configuration at the top of this script.", file=sys.stderr)
        sys.exit(1)

    # Load the docx
    print(f"Loading {test_docx_path}...")
    doc = Document(str(test_docx_path))

    print(f"\nFound {len(doc.paragraphs)} paragraphs")
    print("\n=== Extracting fixture data ===")

    # Extract theme fonts once for the document
    theme_fonts = get_theme_fonts_from_docx_package(doc.part.package)

    # Create a minimal config for annotation extraction
    cfg = UserConfig()
    cfg.display_footnotes = True
    cfg.display_endnotes = True

    # Extract all annotations from the document
    all_comments = get_all_docx_comments(doc)
    all_footnotes = get_all_docx_footnotes(doc, cfg)
    all_endnotes = get_all_docx_endnotes(doc, cfg)

    data = []
    for para_idx, para in enumerate(doc.paragraphs):
        # Extract paragraph-level font properties (from style)
        para_font = {}
        if para.style and para.style.font:
            # Get effective font name (resolves theme fonts)
            effective_font_name = get_effective_font_name_docx(para.style, theme_fonts)

            para_font = {
                "bold": para.style.font.bold,
                "italic": para.style.font.italic,
                "underline": para.style.font.underline,
                "font_name": effective_font_name,  # Use resolved font name
                "font_size": para.style.font.size.pt if para.style.font.size else None,
            }
            # Add color if present
            if para.style.font.color and para.style.font.color.rgb:
                para_font["color_rgb"] = rgb_to_hex(para.style.font.color.rgb)

            # Filter out None values but keep False
            para_font = filter_none_keep_false(para_font)

        para_data = {
            "paragraph_number": para_idx,
            "text": para.text,
            "style": para.style.name if para.style else None,
            "paragraph_font": para_font,
            "runs": []
        }

        for run_idx, run in enumerate(para.runs):
            run_data = {
                "run_number": run_idx,
                "text": run.text,
                "bold": run.font.bold,
                "italic": run.font.italic,
                "underline": run.font.underline,
                "font_name": run.font.name,
                "font_size": run.font.size.pt if run.font.size else None,
            }

            # Add color if present
            if run.font.color and run.font.color.rgb:
                run_data["color_rgb"] = rgb_to_hex(run.font.color.rgb)

            # Extract experimental formatting (exposed by python-docx)
            if run.font.highlight_color:
                run_data["highlight_color"] = run.font.highlight_color.name

            if run.font.strike:
                run_data["strike"] = run.font.strike

            if run.font.double_strike:
                run_data["double_strike"] = run.font.double_strike

            if run.font.subscript:
                run_data["subscript"] = run.font.subscript

            if run.font.superscript:
                run_data["superscript"] = run.font.superscript

            if run.font.all_caps:
                run_data["all_caps"] = run.font.all_caps

            if run.font.small_caps:
                run_data["small_caps"] = run.font.small_caps

            # Extract annotation references from run XML
            try:
                run_xml = run.element.xml
                root = ET.fromstring(run_xml)

                # Find comment references
                comment_refs = root.findall(".//w:commentReference", XML_NS)
                if comment_refs:
                    run_data["comment_refs"] = []
                    for ref in comment_refs:
                        comment_id = ref.get(f'{{{XML_NS["w"]}}}id')
                        if comment_id and comment_id in all_comments:
                            comment_obj = all_comments[comment_id]
                            run_data["comment_refs"].append({
                                "id": comment_id,
                                "text": comment_obj.text,
                                "author": comment_obj.author if hasattr(comment_obj, 'author') else None,
                            })

                # Find footnote references
                footnote_refs = root.findall(".//w:footnoteReference", XML_NS)
                if footnote_refs:
                    run_data["footnote_refs"] = []
                    for ref in footnote_refs:
                        footnote_id = ref.get(f'{{{XML_NS["w"]}}}id')
                        if footnote_id and footnote_id in all_footnotes:
                            footnote_obj = all_footnotes[footnote_id]
                            run_data["footnote_refs"].append({
                                "id": footnote_id,
                                "text_body": footnote_obj.text_body,
                                "hyperlinks": footnote_obj.hyperlinks,
                            })

                # Find endnote references
                endnote_refs = root.findall(".//w:endnoteReference", XML_NS)
                if endnote_refs:
                    run_data["endnote_refs"] = []
                    for ref in endnote_refs:
                        endnote_id = ref.get(f'{{{XML_NS["w"]}}}id')
                        if endnote_id and endnote_id in all_endnotes:
                            endnote_obj = all_endnotes[endnote_id]
                            run_data["endnote_refs"].append({
                                "id": endnote_id,
                                "text_body": endnote_obj.text_body,
                                "hyperlinks": endnote_obj.hyperlinks,
                            })
            except (AttributeError, ET.ParseError):
                pass  # No annotations in this run

            # Filter out None values but keep False
            run_data = filter_none_keep_false(run_data)
            para_data["runs"].append(run_data)

        data.append(para_data)

    # Show what we got
    safe_pprint(data, width=120, item_type="paragraphs")

    # Save to JSON
    output_path = Path(OUTPUT_JSON)
    try:
        output_path.parent.mkdir(parents=True, exist_ok=True)
        with open(output_path, "w", encoding="utf-8") as f:
            json.dump(data, f, indent=2, ensure_ascii=False)
        print(f"\nSaved to {output_path}")
    except (OSError, PermissionError) as e:
        print(f"Error: Failed to write output file: {e}", file=sys.stderr)
        sys.exit(1)


if __name__ == "__main__":
    try:
        main()
    except Exception as e:
        print(f"Fatal error: {e}", file=sys.stderr)
        sys.exit(1)
