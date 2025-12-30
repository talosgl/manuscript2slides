"""Wide scoped tests that touch many parts of the pipeline and program to ensure no catastrophic failures during golden path runs."""

# tests/test_integration.py
from pathlib import Path

import pytest
from docx import Document
from pptx import Presentation  # pyright: ignore[reportPrivateImportUsage]

from manuscript2slides.internals.define_config import PipelineDirection, UserConfig
from manuscript2slides.orchestrator import run_pipeline, run_roundtrip_test


# region test_round_trip
def test_round_trip_preserves_basic_plaintext_content(
    path_to_sample_docx_with_everything: Path, path_to_empty_pptx: Path, tmp_path: Path
) -> None:
    """Integration test: docx -> pptx -> docx preserves content"""

    # Arrange:
    # Get the original docx's text to test against later
    original_docx = Document(str(path_to_sample_docx_with_everything))
    original_text = "\n".join(
        p.text for p in original_docx.paragraphs if p.text.strip()
    )

    # Make a config to send to round_trip
    test_cfg = UserConfig(
        input_docx=path_to_sample_docx_with_everything,
        template_pptx=path_to_empty_pptx,
        output_folder=tmp_path,
    ).enable_all_options()

    _, _, final_docx = run_roundtrip_test(test_cfg)

    final_doc = Document(str(final_docx))
    final_text = "\n".join(p.text for p in final_doc.paragraphs if p.text.strip())

    # Basic check: did we keep most of the text?
    assert len(final_text) > (len(original_text) * 0.8), "Lost too much content"


def test_round_trip_preserves_comments(
    path_to_sample_docx_with_everything: Path, path_to_empty_pptx: Path, tmp_path: Path
) -> None:
    """Integration test: docx -> pptx -> docx preserves comments"""

    # Arrange:
    original_docx = Document(str(path_to_sample_docx_with_everything))

    # Make a config to send to round_trip
    test_cfg = UserConfig(
        input_docx=path_to_sample_docx_with_everything,
        template_pptx=path_to_empty_pptx,
        output_folder=tmp_path,
    ).enable_all_options()

    _, _, final_docx = run_roundtrip_test(test_cfg)

    final_doc = Document(str(final_docx))

    # Check if we preserved comments. We check to see if there is the same or MORE
    # because annotations like footnotes and endnotes get converted to comments in the
    # roundtrip flow.
    assert len(final_doc.comments) >= len(original_docx.comments)


# endregion


# region test_docx2pptx
def test_docx2pptx_creates_valid_output(
    path_to_sample_docx_with_everything: Path, tmp_path: Path, path_to_empty_pptx: Path
) -> None:
    """Integration: docx -> pptx works"""

    # Arrange: Make a config and set the input doc to our tests/data copy of sample_doc.docx
    test_cfg = UserConfig(
        input_docx=path_to_sample_docx_with_everything,
        template_pptx=path_to_empty_pptx,
        output_folder=tmp_path,
    ).enable_all_options()

    # Verify the direction has been set correctly
    assert test_cfg.direction == PipelineDirection.DOCX_TO_PPTX

    # Action: Run the pipeline
    output_path = run_pipeline(test_cfg)

    # Assert: output generated/exists, has the right extension, has slides in it.
    assert output_path.exists()
    assert output_path.suffix == ".pptx"
    pres = Presentation(output_path)
    assert len(pres.slides) > 0


def test_pipeline_fails_gracefully_on_missing_input(
    path_to_empty_pptx: Path, caplog: pytest.LogCaptureFixture
) -> None:
    """Test that if a bad path is passed into pipeline function, it raises without crashing/Traceback."""

    test_cfg = UserConfig(
        input_docx=Path("bad_path.docx"), template_pptx=path_to_empty_pptx
    )  # ensure the template won't be the failure cause for this test

    with pytest.raises(FileNotFoundError):
        run_pipeline(test_cfg)

    assert "ERROR" in caplog.text


def test_pipeline_fails_gracefully_on_invalid_config() -> None:
    """Test we catch invalid config caught before pipeline runs"""
    test_cfg = UserConfig(input_docx=None)

    with pytest.raises(ValueError):
        run_pipeline(test_cfg)


# endregion


# region test_pptx2docx
def test_pptx2docx_creates_valid_output(
    path_to_sample_pptx_with_formatting: Path, tmp_path: Path, path_to_empty_docx: Path
) -> None:
    """Integration: pptx -> docx works"""
    test_cfg = UserConfig(
        input_pptx=path_to_sample_pptx_with_formatting,
        output_folder=tmp_path,
        template_docx=path_to_empty_docx,
    ).enable_all_options()

    assert test_cfg.direction == PipelineDirection.PPTX_TO_DOCX

    output_path = run_pipeline(test_cfg)

    assert output_path.exists()
    assert output_path.suffix == ".docx"
    docu = Document(str(output_path))
    assert len(docu.paragraphs) > 0
