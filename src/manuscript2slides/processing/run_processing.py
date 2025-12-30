# run_processing.py
"""Processes inner-paragraph contents (runs, hyperlinks) for both pipeline directions."""

# For incomplete type stubs in python-pptx:
# pyright: reportAttributeAccessIssue=false
# mypy: disable-error-code="import-untyped"

# region imports
import logging
from typing import Any

from docx import document
from docx.opc import constants
from docx.oxml.ns import qn
from docx.oxml.parser import OxmlElement as OxmlElement_docx
from docx.shared import RGBColor as RGBColor_docx
from docx.text.hyperlink import Hyperlink as Hyperlink_docx
from docx.text.paragraph import Paragraph as Paragraph_docx
from docx.text.run import Run as Run_docx
from pptx.text.text import _Paragraph as Paragraph_pptx
from pptx.text.text import _Run as Run_pptx

from manuscript2slides.annotations.restore_from_slides import (
    safely_extract_comment_data,
    safely_extract_experimental_formatting_data,
)
from manuscript2slides.internals.define_config import UserConfig
from manuscript2slides.models import SlideNotes
from manuscript2slides.processing.docx_xml import detect_field_code_hyperlinks
from manuscript2slides.processing.formatting import (
    apply_experimental_formatting_from_metadata,
    copy_paragraph_formatting_docx2pptx,
    copy_run_formatting_docx2pptx,
    copy_run_formatting_pptx2docx,
)

# endregion

log = logging.getLogger("manuscript2slides")


# region process_docx_paragraph_inner_contents
def process_docx_paragraph_inner_contents(
    paragraph: Paragraph_docx, pptx_paragraph: Paragraph_pptx, cfg: UserConfig
) -> list[dict]:
    """
    Iterate through a paragraph's runs and hyperlinks, in document order, and:
    - copy it into the slide, with formatting
    - capture any additional metadata (like experimental formatting) that we cannot apply directly to the copied-over
        pptx objects
    """
    items_processed = False
    experimental_formatting_metadata: list[dict[str, Any]] = []

    # Copy baseline paragraph-level font and basic formatting before processing runs
    if paragraph.style and paragraph.style.font:
        copy_paragraph_formatting_docx2pptx(paragraph, pptx_paragraph)

    for item in paragraph.iter_inner_content():
        items_processed = True
        if isinstance(item, Run_docx):
            # If this Run has a field code for instrText and it begins with HYPERLINK, this is an old-style
            # word hyperlink, which we cannot handle the same way as normal docx hyperlinks. But we try to detect
            # when it happens and report it to the user.
            field_code_URL = detect_field_code_hyperlinks(item)
            if field_code_URL:
                item.text = f" [Link: {field_code_URL}] "

            process_docx_run(
                item, pptx_paragraph, experimental_formatting_metadata, cfg
            )

        # Run and Hyperlink objects are peers in docx, but Hyperlinks can contain lists of Runs.
        # We check the item.url field because that seems the most reliable way to see if this is a
        # basic run versus a Hyperlink containing its own nested runs.
        # https://python-docx.readthedocs.io/en/latest/api/text.html#docx.text.hyperlink.Hyperlink.url
        elif isinstance(item, Hyperlink_docx):
            # Process all runs within the Hyperlink
            for run in item.runs:
                process_docx_run(
                    run,
                    pptx_paragraph,
                    experimental_formatting_metadata,
                    cfg,
                    item,
                )
        else:
            log.warning(f"Unknown content type in paragraph: {type(item)}")

    # Fallback: if no content was processed but paragraph has text
    if not items_processed and paragraph.text:
        log.debug(
            f"Fallback: paragraph has text but no runs/hyperlinks: {paragraph.text[:50]}"
        )
        pptx_run = pptx_paragraph.add_run()
        pptx_run.text = paragraph.text

    return experimental_formatting_metadata


# endregion


# region process_docx_run
def process_docx_run(
    run: Run_docx,
    pptx_paragraph: Paragraph_pptx,
    experimental_formatting_metadata: list,
    cfg: UserConfig,
    hyperlink: Hyperlink_docx | None = None,
) -> Run_pptx:
    """Copy a run from the docx parent to the pptx paragraph, including copying its formatting."""
    # Handle formatting

    pptx_run = pptx_paragraph.add_run()
    copy_run_formatting_docx2pptx(run, pptx_run, experimental_formatting_metadata, cfg)

    if hyperlink:
        pptx_run_url = pptx_run.hyperlink

        # Only external hyperlinks are supported in v1
        if hyperlink.url:
            pptx_run_url.address = hyperlink.url

        # Future: handle hyperlink.fragment (document anchors) here # TODO, leafy
        # elif hyperlink.fragment:
        #   # Look up slide and create internal link
        #   # (Also, will need to handle restoring the anchor in pptx2docx pipeline
        #   # via either storing metadata here, or another intelligent parse/conversion
        #   # from the data in the slide)

    return pptx_run


# endregion


# region process_pptx_run
def process_pptx_run(
    run: Run_pptx,
    new_para: Paragraph_docx,
    new_doc: document.Document,
    slide_notes: SlideNotes,
    matched_comment_ids: set,
    cfg: UserConfig,
) -> Run_docx:
    """
    Process a single run from a pptx slide paragraph by copying its basic formatting into a new docx run, and detecting if its text content
    matches experimental formatting metadata, and/or comment metadata from the speaker notes JSON.
    """

    # Handle adding hyperlinks versus regular runs, and the runs' basic formatting.
    if run.hyperlink.address:
        log.debug(f"Hyperlink address found: {run.hyperlink.address}")
        run_from_hyperlink = add_hyperlink_to_docx_paragraph(
            new_para, run.hyperlink.address, run.text
        )
        last_run = run_from_hyperlink
        copy_run_formatting_pptx2docx(run, run_from_hyperlink, cfg)
    else:
        last_run = new_para.add_run()
        copy_run_formatting_pptx2docx(run, last_run, cfg)

    # Check if this run contains matching text for comments from the this slide's speaker notes' stored JSON
    # metadata, from previous docx2pptx pipeline processing
    if slide_notes.comments:
        # Check to see if the run text matches any comments' ref text
        for comment in slide_notes.comments:
            comment_data = safely_extract_comment_data(comment)

            if comment_data is None:
                log.debug(f"Skipping invalid comment: {comment}")
                continue

            if (
                comment_data["reference_text"] in run.text
                and comment_data["id"] not in matched_comment_ids
            ):
                new_doc.add_comment(
                    last_run,
                    comment_data["text"],
                    comment_data["author"],
                    comment_data["initials"],
                )
                matched_comment_ids.add(comment_data["id"])
            # don't break; there can be multiple comments added to a single run

    if cfg.experimental_formatting_on and slide_notes.experimental_formatting:
        for exp_fmt in slide_notes.experimental_formatting:
            fmt_info = safely_extract_experimental_formatting_data(exp_fmt)
            if fmt_info is None:
                continue
            if exp_fmt["ref_text"] in run.text:
                apply_experimental_formatting_from_metadata(last_run, exp_fmt)

    return last_run


# endregion


def add_hyperlink_to_docx_paragraph(
    paragraph: Paragraph_docx, url: str, text: str
) -> Run_docx:
    """
    Custom function to add Hyperlink objects to docx paragraphs using XML manipulation.

    - Create a regular run using paragraph.add_run()
    - Create the hyperlink XML element structure
    - Move the run's XML element from being a direct child of the paragraph to being a nested child of the Hyperlink element
    - Add the Hyperlink element (which now contains the run) to the paragraph

    Adapted from/referenced https://stackoverflow.com/questions/47666642/adding-an-hyperlink-in-msword-by-using-python-docx
    and https://github.com/python-openxml/python-docx/issues/384#issuecomment-294853130
    """
    try:
        # Step 1: Create the hyperlink structure
        part = paragraph.part
        r_id = part.relate_to(
            url, constants.RELATIONSHIP_TYPE.HYPERLINK, is_external=True
        )
        hyperlink = OxmlElement_docx("w:hyperlink")
        hyperlink.set(qn("r:id"), r_id)

        # 2. Create the Run with PARAGRAPH as the parent (to avoid an XML .part error)
        # Use a new 'w:r' element
        new_run = Run_docx(OxmlElement_docx("w:r"), paragraph)
        new_run.text = text

        # 3. Try to apply formatting (manual blue/underline)
        try:
            new_run.font.color.rgb = RGBColor_docx(5, 99, 193)
            new_run.font.underline = True
        except Exception as e:
            log.warning(f"Could not apply hyperlink formatting: {e}")
            # Continue anyway - link will work, just won't be blue/underlined

        # 4. NOW move the run's internal XML inside the hyperlink XML
        hyperlink.append(new_run._element)

        # 5. Append the whole hyperlink structure to the paragraph XML
        paragraph._p.append(hyperlink)

        return new_run

    except Exception as e:
        # Fallback: add as plaintext
        log.warning(
            f"Could not create hyperlink for '{text}' (url: {url}): {e}. Adding as plain text."
        )
        return paragraph.add_run(text)


# endregion
