# chunking.py
"""Create docx2pptx chunks including by paragraph, page, heading (flat), and heading (nested).
All strategies break a new chunk on a docx page break to prevent slide text overflow."""

import logging

from docx import document

from manuscript2slides.internals.define_config import ChunkType
from manuscript2slides.internals.run_context import get_pipeline_run_id
from manuscript2slides.models import Chunk_docx

log = logging.getLogger("manuscript2slides")


# region create_docx_chunks Orchestrator
def create_docx_chunks(
    doc: document.Document, chunk_type: ChunkType
) -> list[Chunk_docx]:
    """
    Orchestrator function to create chunks (that will become slides) from the document
    contents, either from paragraph, heading (heading_nested or heading_flat),
    or page. Defaults to paragraph.
    """
    if chunk_type == ChunkType.HEADING_FLAT:
        chunks = chunk_by_heading_flat(doc)
    elif chunk_type == ChunkType.HEADING_NESTED:
        chunks = chunk_by_heading_nested(doc)
    elif chunk_type == ChunkType.PAGE:
        chunks = chunk_by_page(doc)
    else:
        chunks = chunk_by_paragraph(doc)
    return chunks


# endregion


# region by Paragraph
def chunk_by_paragraph(doc: document.Document) -> list[Chunk_docx]:
    """
    Creates chunks (which will become slides) based on paragraph, which are blocks of content
    separated by whitespace.
    """
    pipeline_id = get_pipeline_run_id()

    log.info(f"Running the chunk by paragraph strategy. [pipeline:{pipeline_id}]")

    paragraph_chunks: list[Chunk_docx] = []

    # Start at page 1
    current_page_number = 1

    for para in doc.paragraphs:
        if para.contains_page_break:
            # Increment page count
            current_page_number = current_page_number + 1

        # Skip empty paragraphs (but keep those that are new-lines to respect intentional whitespace newlines)
        if para.text == "":
            log.debug("Skipping empty paragraph.")
            continue

        log.debug(f"Paragraph begins: {para.text[:30]}... [pipeline:{pipeline_id}]")

        new_chunk = Chunk_docx.create_with_paragraph(para)
        new_chunk.original_sequence_number = current_page_number
        paragraph_chunks.append(new_chunk)

    log.info(
        f"We processed {len(paragraph_chunks)} paragraph chunks. [pipeline:{pipeline_id}]"
    )

    return paragraph_chunks


# endregion


# region by Page
def chunk_by_page(doc: document.Document) -> list[Chunk_docx]:
    """Creates chunks based on page breaks"""
    pipeline_id = get_pipeline_run_id()

    log.info(f"Running the chunk by page strategy. [pipeline:{pipeline_id}]")

    # Start building the chunks
    all_chunks: list[Chunk_docx] = []

    # Start at page 1
    current_page_number = 1

    # Start with a current chunk ready-to-go
    current_page_chunk: Chunk_docx = Chunk_docx(
        original_sequence_number=current_page_number
    )

    for para in doc.paragraphs:
        if para.contains_page_break:
            # Increment page count
            current_page_number = current_page_number + 1

        # Skip empty paragraphs (keep intentional whitespace newlines)
        if para.text == "":
            log.debug("Skipping empty paragraph.")
            continue

        log.debug(f"Paragraph begins: {para.text[:30]}... [pipeline:{pipeline_id}]")

        # If the current_page_chunk is empty, append the current para regardless of style & continue to next para.
        if not current_page_chunk.paragraphs:
            current_page_chunk.add_paragraph(para)
            continue

        # Handle page breaks - create new chunk and start fresh
        if para.contains_page_break:
            # Add the current_chunk to chunks list (if it has content)
            if current_page_chunk:
                all_chunks.append(current_page_chunk)

            # Start new chunk with this paragraph
            current_page_chunk = Chunk_docx.create_with_paragraph(para)
            current_page_chunk.original_sequence_number = current_page_number

            continue

        # If there was no page break, just append this paragraph to the current_chunk
        current_page_chunk.add_paragraph(para)

    # Ensure final chunk from loop is added to chunks list
    if current_page_chunk:
        all_chunks.append(current_page_chunk)

    log.info(f"We processed {len(all_chunks)} page chunks. [pipeline:{pipeline_id}]")
    return all_chunks


# endregion


# region by Heading (nested)
def chunk_by_heading_nested(doc: document.Document) -> list[Chunk_docx]:
    """
    Creates chunks based on headings, using nesting logic to group "deeper" headings

    New chunks are begun when:
    - a page break happens in the middle of a paragraph
    - we reach a heading-depth that is equal to or "higher" than (number is less than) the current depth
    Otherwise, appends paragraph to the current chunk.

    E.g.:

    CHUNK:
    H1
    Normal Paragraph
    H2
    Normal Paragraph
    Normal Paragraph

    NEXT_CHUNK:
    H2
    Normal Paragraph
    Normal Paragraph
    Normal Paragraph

    NEXT_CHUNK:
    H1
    H2
    Pararaph
    Normal Paragraph
    H3
    Normal Paragraph

    NEXT_CHUNK:
    H2
    Normal Paragraph
    Normal Paragraph

    """
    pipeline_id = get_pipeline_run_id()
    log.info(
        f"Running the chunk by heading (nested) strategy. [pipeline:{pipeline_id}]"
    )
    # Start at page 1
    current_page_number = 1

    # Start building the chunks
    all_chunks: list[Chunk_docx] = []
    current_chunk: Chunk_docx = Chunk_docx(original_sequence_number=current_page_number)

    # Initialize current_heading_style_name
    current_heading_style_name = "Normal"  # Default for documents without headings

    for i, para in enumerate(doc.paragraphs):
        if para.contains_page_break:
            # Increment page number
            current_page_number = current_page_number + 1

        # Skip empty paragraphs
        if para.text == "":
            log.debug("Skipping empty paragraph.")
            continue

        # Set a style_name to make Pylance happy (it gets mad if we direct-check para.style.style_name later)
        style_name = para.style.name if para.style and para.style.name else "Normal"

        log.debug(
            f"Paragraph begins: {para.text[:30]}... and is index: {i}. [pipeline:{pipeline_id}]"
        )

        # If the current_chunk is empty, append the current para regardless of style & continue to next para.
        if not current_chunk.paragraphs:
            current_chunk.add_paragraph(para)
            if is_standard_heading(style_name):
                current_heading_style_name = style_name
            continue

        # Handle page breaks - create new chunk and start fresh
        if para.contains_page_break:
            # Add the current chunk to chunks list (if it has content)
            if current_chunk:
                all_chunks.append(current_chunk)

            # Start new chunk with this paragraph
            current_chunk = Chunk_docx.create_with_paragraph(para)
            current_chunk.original_sequence_number = current_page_number

            # Update heading depth if this paragraph is a heading
            if is_standard_heading(style_name):
                current_heading_style_name = style_name
            continue

        # Handle headings
        if is_standard_heading(style_name):
            # Check if this heading is at same level or higher (less deep) than current. Smaller numbers are higher up in the hierarchy.
            if get_heading_level(style_name) <= get_heading_level(
                current_heading_style_name
            ):
                # If yes, start a new chunk
                if current_chunk:
                    all_chunks.append(current_chunk)
                current_chunk = Chunk_docx.create_with_paragraph(para)
                current_chunk.original_sequence_number = current_page_number
                current_heading_style_name = style_name
            else:
                # This heading is deeper, add to current chunk
                current_chunk.add_paragraph(para)
                current_heading_style_name = style_name
        else:
            # Normal paragraph - add to current chunk
            current_chunk.add_paragraph(para)

    if current_chunk:
        all_chunks.append(current_chunk)

    log.info(
        f"We processed {len(all_chunks)} nested heading chunks. [pipeline:{pipeline_id}]"
    )
    return all_chunks


# endregion


# region by Heading (flat)
def chunk_by_heading_flat(doc: document.Document) -> list[Chunk_docx]:
    """
    Creates chunks based on headings; no nesting logic used. New chunks are created when:
    - a page break happens in the middle of a paragraph
    - we reach any paragraph that is any heading

    CHUNK:
    H1
    Normal Paragraph

    NEXT_CHUNK:
    H2
    Normal Paragraph
    Normal Paragraph

    NEXT_CHUNK:
    H2
    Normal Paragraph
    Normal Paragraph
    Normal Paragraph

    NEXT_CHUNK:
    H1

    NEXT_CHUNK:
    H2
    Pararaph
    Normal Paragraph

    NEXT_CHUNK:
    H3
    Normal Paragraph

    NEXT_CHUNK:
    H2
    Normal Paragraph
    Normal Paragraph
    """
    pipeline_id = get_pipeline_run_id()
    log.info(f"Running the chunk by heading (flat) strategy. [pipeline:{pipeline_id}]")

    # Start at page 1
    current_page_number = 1

    # Start building the chunks
    all_chunks: list[Chunk_docx] = []
    current_chunk: Chunk_docx = Chunk_docx(original_sequence_number=current_page_number)

    for para in doc.paragraphs:
        if para.contains_page_break:
            # Increment page count
            current_page_number = current_page_number + 1

        # Skip empty paragraphs
        if para.text == "":
            log.debug("Skipping empty paragraph.")
            continue

        # Set a style_name to make Pylance happy (it gets mad if we direct-check para.style.name later)
        style_name = para.style.name if para.style and para.style.name else "Normal"

        log.debug(f"Paragraph begins: {para.text[:30]}... [pipeline:{pipeline_id}]")

        # If the current_chunk is empty, append the current para regardless of style & continue to next para.
        if not current_chunk.paragraphs:
            current_chunk.add_paragraph(para)
            continue

        # Handle page breaks - always start a new chunk
        if para.contains_page_break:
            # Add the current chunk to chunks list (if it has content)
            if current_chunk:
                all_chunks.append(current_chunk)

            # Start new chunk with this paragraph
            current_chunk = Chunk_docx.create_with_paragraph(para)
            current_chunk.original_sequence_number = current_page_number
            continue

        # If this paragraph is a heading, start a new chunk
        if is_standard_heading(style_name):
            # If we already have content in current_chunk, save it and start fresh
            if current_chunk:
                all_chunks.append(current_chunk)

            # Start new chunk with this paragraph
            current_chunk = Chunk_docx.create_with_paragraph(para)
            current_chunk.original_sequence_number = current_page_number

        else:
            # This is a normal paragraph - add it to current chunk
            current_chunk.add_paragraph(para)

    if current_chunk:
        all_chunks.append(current_chunk)

    log.info(
        f"We processed {len(all_chunks)} flat heading chunks. [pipeline:{pipeline_id}]"
    )
    return all_chunks


# endregion


# region heading helpers


# region _is_standard_heading
def is_standard_heading(style_name: str) -> bool:
    """Check if paragraph.style.name is a standard Word Heading (Heading 1, Heading 2, ..., Heading 6)"""
    return style_name.startswith("Heading") and style_name[8:].isdigit()


# endregion


# region get_heading_level
def get_heading_level(style_name: str) -> int | float:
    """
    Extract the numeric level from a heading style name (e.g., 'Heading 2' -> 2),
    or return infinity if the style name doesn't have a number.
    """
    try:
        return int(style_name[8:])
    except (ValueError, IndexError):
        return float("inf")  # Treat non-headings as "deepest possible"


# endregion

# endregion
