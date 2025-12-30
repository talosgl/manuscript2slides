# models.py
"""Data models for document chunks and annotations."""

from __future__ import annotations

from dataclasses import dataclass, field

from docx.comments import Comment as Comment_docx
from docx.text.paragraph import Paragraph as Paragraph_docx


# region Docx Annotation custom classes
# region Comment_docx_custom
@dataclass
class Comment_docx_custom:
    """A custom wrapper for the python-docx Comment class, allowing us to capture reference text."""

    comment_obj: Comment_docx
    reference_text: str | None = None  # The text this comment is attached to

    @property
    def note_id(self) -> int:
        """Alias for comment_id to provide a common interface with other note types."""
        return self.comment_obj.comment_id


# endregion


# region Footnote_docx
@dataclass
class Footnote_docx:
    """
    Represents a footnote extracted from a docx.

    Contains the footnote ID, text content, and any hyperlinks found within.
    Used for preserving footnote information when python-docx doesn't provide
    direct access to footnote content.
    """

    footnote_id: str
    text_body: str
    hyperlinks: list[str] = field(default_factory=list)
    reference_text: str | None = None

    @property
    def note_id(self) -> str:
        """Alias for footnote_id to provide a common interface with other note types."""
        return self.footnote_id


# endregion


# region Endnote_docx
@dataclass
class Endnote_docx:
    """
    Represents a endnote extracted from a docx.

    Contains the endnote ID, text content, and any hyperlinks found within.
    Used for preserving endnote information when python-docx doesn't provide
    direct access to endnote content.
    """

    endnote_id: str
    text_body: str
    hyperlinks: list[str] = field(default_factory=list)
    reference_text: str | None = None

    @property
    def note_id(self) -> str:
        """Alias for endnote_id to provide a common interface with other note types."""
        return self.endnote_id


# endregion
# endregion


# region Chunk_docx
@dataclass
class Chunk_docx:
    """Class for Chunk objects made from docx paragraphs and their associated annotations."""

    from typing import Optional

    # Page or slide where this chunk came from
    original_sequence_number: int = 0

    # Use "default_factory" to ensure every chunk gets its own list.
    # (Lists are mutable; it is a common error/bug to accidentally assign one list
    # shared amongst every instance of a class, rather than one per instance.)
    paragraphs: list[Paragraph_docx] = field(default_factory=list[Paragraph_docx])

    comments: list[Comment_docx_custom] = field(
        default_factory=list[Comment_docx_custom]
    )
    footnotes: list[Footnote_docx] = field(default_factory=list[Footnote_docx])
    endnotes: list[Endnote_docx] = field(default_factory=list[Endnote_docx])

    @classmethod
    def create_with_paragraph(cls, paragraph: Paragraph_docx) -> "Chunk_docx":
        """Create a new instance of a Chunk_docx object but also populate the paragraphs list with an initial element."""
        return cls(paragraphs=[paragraph])

    def add_paragraph(self, new_paragraph: Paragraph_docx) -> None:
        """Add a paragraph to this Chunk object's paragraphs list."""
        self.paragraphs.append(new_paragraph)

    def add_paragraphs(self, new_paragraphs: list[Paragraph_docx]) -> None:
        """Add a list of paragraphs to this Chunk object's paragraphs list."""
        self.paragraphs.extend(new_paragraphs)  # Add multiple at once

    def add_comment(self, comment: Comment_docx_custom) -> None:
        """Add a comment to this Chunk object's comment list."""
        self.comments.append(comment)

    def add_footnote(self, footnote: Footnote_docx) -> None:
        """Add a footnote to this Chunk object's footnote list."""
        self.footnotes.append(footnote)

    def add_endnote(self, endnote: Endnote_docx) -> None:
        """Add a endnote to this Chunk object's endnote list."""
        self.endnotes.append(endnote)


# endregion


# region SlideNotes
@dataclass
class SlideNotes:
    """User notes and metadata extracted from a slide's speaker notes."""

    metadata: dict = field(default_factory=dict)
    user_notes: str = ""
    comments: list = field(default_factory=list)
    footnotes: list = field(default_factory=list)
    endnotes: list = field(default_factory=list)
    headings: list = field(default_factory=list)
    experimental_formatting: list = field(default_factory=list)

    @property
    def has_metadata(self) -> bool:
        """Returns a bool to indicate whether we did or did not find/store JSON metadata from these SlideNotes."""
        return bool(self.metadata)  # True if dict is non-empty

    @property
    def has_user_notes(self) -> bool:
        """
        Returns a bool to indicate whether we did or did not find/store unique user notes (not JSON metadata, and not
        copied annotations from earlier docx2pptx pipeline runs) from these SlideNotes.
        """
        return bool(self.user_notes.strip())


# endregion
